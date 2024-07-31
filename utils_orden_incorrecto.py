import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import json
import os
import sys
import re

def get_config_file_path():
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, 'styles_config.json')
    else:
        return os.path.join(os.path.dirname(__file__), 'styles_config.json')

def load_config():
    config_file = get_config_file_path()
    print(f"Intentando cargar configuración desde {config_file}")
    if os.path.exists(config_file):
        try:
            with open(config_file, 'r') as f:
                config = json.load(f)
                print(f"Configuración cargada: {config}")
                return config
        except json.JSONDecodeError as e:
            print(f"Error al decodificar JSON: {e}")
            return {}
    else:
        print(f"Archivo de configuración no encontrado: {config_file}")
        return {}

def save_config(config):
    config_file = get_config_file_path()
    try:
        with open(config_file, 'w') as f:
            json.dump(config, f, indent=4)
        print(f"Configuración guardada en {config_file}")
    except Exception as e:
        print(f"Error al guardar la configuración: {e}")

def apply_styles(paragraph, text, style_name, style_type, doc):
    styles = doc.styles
    if style_name not in [style.name for style in styles]:
        if style_type == 'parrafo':
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        elif style_type == 'caracter':
            style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
        font = style.font
        font.size = Pt(12)

    if style_type == 'parrafo':
        paragraph.style = style_name
        paragraph.add_run(text)
    elif style_type == 'caracter':
        run = paragraph.add_run(text)
        run.style = style_name
    print(f"Applied style {style_name} with text: {text}")

def clean_default_styles(doc, config):
    styles = doc.styles
    keep_styles = {s['style'] for s in config.values()}
    keep_styles.add("Agenda-General-Parrafo")

    for style in list(styles):
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER) and style.name not in keep_styles:
            styles.element.remove(style.element)

def sanitize_xml_content(content):
    content = re.sub(r'&(?!amp;)', 'i', content)
    replacements = {}
    for old, new in replacements.items():
        content = content.replace(old, new)
    return content

def create_sanitized_copy(xml_file):
    sanitized_file = xml_file + ".sanitized"
    with open(xml_file, 'r', encoding='utf-8') as file:
        content = file.read()
    sanitized_content = sanitize_xml_content(content)
    with open(sanitized_file, 'w', encoding='utf-8') as file:
        file.write(sanitized_content)
    return sanitized_file

def validate_xml_file(xml_file):
    try:
        tree = ET.parse(xml_file)
        tree.getroot()
        return True
    except ET.ParseError:
        return False

def apply_combined_styles(element1, element2, tag1, tag2, doc, config):
    text1 = element1.text.strip() if element1 is not None and element1.text else ""
    text2 = element2.text.strip() if element2 is not None and element2.text else ""
    combined_text = f"{text1} · {text2}" if text1 and text2 else text1 or text2

    if combined_text:
        paragraph = doc.add_paragraph()
        style_name1 = config.get(tag1, {}).get('style', "Agenda-General-Parrafo")
        style_type1 = config.get(tag1, {}).get('type', 'parrafo')
        apply_styles(paragraph, combined_text, style_name1, style_type1, doc)

def process_fields(element, parent_tag="", doc=None, config=None):
    for child in element:
        tag = f"{parent_tag}/{child.tag}" if parent_tag else child.tag
        if child.tag in ["Evento-Principal-Programa", "Sub-evento-descripcion", "Sub-evento-actividades"]:
            process_fields(child, tag, doc=doc, config=config)
        elif child.tag.startswith("actividad"):
            for actividad in child.findall("actividad"):
                process_fields(actividad, f"{tag}/actividad", doc=doc, config=config)
        elif child.text and child.text.strip():
            paragraph = doc.add_paragraph()
            style_name = config.get(tag, {}).get('style', "Agenda-General-Parrafo")
            style_type = config.get(tag, {}).get('type', 'parrafo')
            apply_styles(paragraph, child.text.strip(), style_name, style_type, doc)
        elif "Hora" in child.tag or "Lugar" in child.tag:
            if "Hora" in child.tag:
                lugar_tag = tag.replace("Hora", "Lugar")
                lugar_element = element.find(lugar_tag)
                apply_combined_styles(child, lugar_element, tag, lugar_tag, doc, config)
            elif "Lugar" in child.tag:
                hora_tag = tag.replace("Lugar", "Hora")
                hora_element = element.find(hora_tag)
                apply_combined_styles(hora_element, child, hora_tag, tag, doc, config)

def process_xml_to_docx(xml_file, output_folder, output_file_name):
    print("Iniciando procesamiento del archivo XML.")
    sanitized_file = create_sanitized_copy(xml_file)
    if not validate_xml_file(sanitized_file):
        print("Error: El archivo XML no está bien formateado después de la sanitización.")
        return

    config = load_config()
    print("Configuración cargada:", config)
    tree = ET.parse(sanitized_file)
    root = tree.getroot()

    doc = Document()

    if "Agenda-General-Parrafo" not in [style.name for style in doc.styles]:
        general_style = doc.styles.add_style("Agenda-General-Parrafo", WD_STYLE_TYPE.PARAGRAPH)
        general_style.font.size = Pt(12)

    for event in root.findall('Evento-Principal'):
        print(f"Procesando evento: {event.tag}")
        process_fields(event, doc=doc, config=config)

        programa = event.find('Evento-Principal-Programa')
        if programa is not None:
            for sub_event in programa.findall('Sub-evento'):
                print(f"Procesando sub-evento: {sub_event.tag}")
                process_fields(sub_event, doc=doc, config=config)

                actividades = sub_event.find('Sub-evento-actividades')
                if actividades is not None:
                    for actividad in actividades.findall('actividad'):
                        print(f"Procesando actividad: {actividad.tag}")
                        process_fields(actividad, doc=doc, config=config)

    clean_default_styles(doc, config)

    for paragraph in doc.paragraphs:
        if paragraph.style is None or paragraph.style.name == 'Normal':
            paragraph.style = "Agenda-General-Parrafo"

    output_path = os.path.join(output_folder, output_file_name)
    doc.save(output_path)
    print(f"Documento guardado en {output_path}")

    if os.path.exists(sanitized_file):
        os.remove(sanitized_file)
        print(f"Archivo sanitizado {sanitized_file} eliminado.")

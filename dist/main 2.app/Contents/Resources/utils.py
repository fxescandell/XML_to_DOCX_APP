import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import json
import os
import re
import sys

def get_config_file_path():
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, 'styles_config.json')
    else:
        return os.path.join(os.path.dirname(__file__), 'styles_config.json')

def load_config():
    config_file = get_config_file_path()
    print(f"Intentando cargar configuración desde {config_file}")
    if os.path.exists(config_file):
        try:
            with open(config_file, 'r') as f:
                config = json.load(f)
                print(f"Configuración cargada: {config}")
                return config
        except json.JSONDecodeError as e:
            print(f"Error al decodificar JSON: {e}")
            return {}
    else:
        print(f"Archivo de configuración no encontrado: {config_file}")
        return {}

def save_config(config):
    config_file = get_config_file_path()
    try:
        with open(config_file, 'w') as f:
            json.dump(config, f, indent=4)
        print(f"Configuración guardada en {config_file}")
    except Exception as e:
        print(f"Error al guardar la configuración: {e}")

def apply_styles(paragraph, text, style_name, style_type, doc):
    styles = doc.styles
    if style_name not in [style.name for style in styles]:
        if style_type == 'parrafo':
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        elif style_type == 'caracter':
            style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
        font = style.font
        font.size = Pt(12)  # Ajusta el tamaño de fuente si es necesario

    if style_type == 'parrafo':
        paragraph.style = style_name
        paragraph.add_run(text)
    elif style_type == 'caracter':
        run = paragraph.add_run(text)
        run.style = style_name
    print(f"Applied style {style_name} with text: {text}")

def clean_default_styles(doc):
    styles = doc.styles
    keep_styles = {s['style'] for s in load_config().values()}
    keep_styles.add("Agenda-General-Parrafo")  # Incluye el estilo de párrafo general

    for style in list(styles):
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER) and style.name not in keep_styles:
            styles.element.remove(style.element)

def sanitize_xml_content(content):
    # Reemplaza & sólo si no está seguido de amp; para evitar interferir con entidades HTML válidas
    content = re.sub(r'&(?!amp;)', 'i', content)
    replacements = {
        # Si queremos sanitizar otros caracteres, van aquí. Ejemplo: '&': 'i',
    }
    for old, new in replacements.items():
        content = content.replace(old, new)
    return content

def create_sanitized_copy(xml_file):
    sanitized_file = xml_file + ".sanitized"
    with open(xml_file, 'r', encoding='utf-8') as file:
        content = file.read()
    sanitized_content = sanitize_xml_content(content)
    with open(sanitized_file, 'w', encoding='utf-8') as file:
        file.write(sanitized_content)
    return sanitized_file

def validate_xml_file(xml_file):
    try:
        tree = ET.parse(xml_file)
        tree.getroot()
        return True
    except ET.ParseError:
        return False

def process_combined_elements(elements, field1, field2, paragraph, config, doc):
    text1 = elements[0].text.strip() if elements[0] is not None and elements[0].text else ""
    text2 = elements[1].text.strip() if elements[1] is not None and elements[1].text else ""

    print(f"Processing combined elements: {field1}='{text1}', {field2}='{text2}'")

    if text1:
        style_name1 = config.get(field1, {}).get('style', "Agenda-General-Parrafo")
        style_type1 = config.get(field1, {}).get('type', 'caracter')  # Por defecto 'caracter'
        apply_styles(paragraph, text1, style_name1, style_type1, doc)

    if text1 and text2:
        paragraph.add_run(" · ")

    if text2:
        style_name2 = config.get(field2, {}).get('style', "Agenda-General-Parrafo")
        style_type2 = config.get(field2, {}).get('type', 'caracter')  # Por defecto 'caracter'
        apply_styles(paragraph, text2, style_name2, style_type2, doc)

def process_fields(parent_element, fields, doc, config):
    processed_fields = set()
    for field in fields:
        if field in ["Evento-Principal-Hora", "Sub-evento-Hora", "actividad-hora"]:
            if field not in processed_fields:
                paragraph = None
                if field == "Evento-Principal-Hora":
                    if parent_element.find('Evento-Principal-Hora') is not None or parent_element.find('Evento-Principal-Lugar') is not None:
                        paragraph = doc.add_paragraph()
                        process_combined_elements([parent_element.find('Evento-Principal-Hora'), parent_element.find('Evento-Principal-Lugar')], 'Evento-Principal-Hora', 'Evento-Principal-Lugar', paragraph, config, doc)
                        processed_fields.update(['Evento-Principal-Hora', 'Evento-Principal-Lugar'])
                elif field == "Sub-evento-Hora":
                    if parent_element.find('Sub-evento-Hora') is not None or parent_element.find('Sub-evento-Lugar') is not None:
                        paragraph = doc.add_paragraph()
                        process_combined_elements([parent_element.find('Sub-evento-Hora'), parent_element.find('Sub-evento-Lugar')], 'Sub-evento-Hora', 'Sub-evento-Lugar', paragraph, config, doc)
                        processed_fields.update(['Sub-evento-Hora', 'Sub-evento-Lugar'])
                elif field == "actividad-hora":
                    if parent_element.find('actividad-hora') is not None or parent_element.find('actividad-lugar') is not None:
                        paragraph = doc.add_paragraph()
                        process_combined_elements([parent_element.find('actividad-hora'), parent_element.find('actividad-lugar')], 'actividad-hora', 'actividad-lugar', paragraph, config, doc)
                        processed_fields.update(['actividad-hora', 'actividad-lugar'])
        else:
            if field not in processed_fields:
                element = parent_element.find(field)
                if element is not None and element.text and element.text.strip():
                    paragraph = doc.add_paragraph()
                    style_name = config.get(field, {}).get('style', "Agenda-General-Parrafo")
                    style_type = config.get(field, {}).get('type', 'parrafo')  # Por defecto 'parrafo'
                    print(f"Processing field: {field} with text: {element.text.strip()}")
                    apply_styles(paragraph, element.text.strip(), style_name, style_type, doc)
                    processed_fields.add(field)

def process_xml_to_docx(xml_file, output_folder, output_file_name):
    print("Iniciando procesamiento del archivo XML.")
    sanitized_file = create_sanitized_copy(xml_file)
    if not validate_xml_file(sanitized_file):
        print("Error: El archivo XML no está bien formateado después de la sanitización.")
        return

    config = load_config()
    print("Configuración cargada:", config)
    tree = ET.parse(sanitized_file)
    root = tree.getroot()

    doc = Document()

    # Crear el estilo "Agenda-General-Parrafo"
    if "Agenda-General-Parrafo" not in [style.name for style in doc.styles]:
        general_style = doc.styles.add_style("Agenda-General-Parrafo", WD_STYLE_TYPE.PARAGRAPH)
        general_style.font.size = Pt(12)

    for event in root.findall('Evento-Principal'):
        print(f"Procesando evento: {event.tag}")
        process_fields(event, config.keys(), doc, config)

        programa = event.find('Evento-Principal-Programa')
        if programa is not None:
            for sub_event in programa.findall('Sub-evento'):
                print(f"Procesando sub-evento: {sub_event.tag}")
                process_fields(sub_event, config.keys(), doc, config)

                actividades = sub_event.find('Sub-evento-actividades')
                if actividades is not None:
                    for actividad in actividades.findall('actividad'):
                        print(f"Procesando actividad: {actividad.tag}")
                        process_fields(actividad, config.keys(), doc, config)

    # Limpiar los estilos por defecto
    clean_default_styles(doc)

    # Aplicar el estilo "Agenda-General-Parrafo" a todos los párrafos
    for paragraph in doc.paragraphs:
        if paragraph.style is None or paragraph.style.name == 'Normal':  # Cambiar el estilo Normal por defecto a Agenda-General-Parrafo
            paragraph.style = "Agenda-General-Parrafo"

    output_path = os.path.join(output_folder, output_file_name)
    doc.save(output_path)
    print(f"Documento guardado en {output_path}")

    # Eliminar el archivo sanitizado después del procesamiento
    if os.path.exists(sanitized_file):
        os.remove(sanitized_file)
        print(f"Archivo sanitizado {sanitized_file} eliminado.")

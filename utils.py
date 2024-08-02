import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import json
import os
import sys
import re

def get_config_file_path():
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, 'styles_config.json')
    else:
        return os.path.join(os.path.dirname(__file__), 'styles_config.json')

def load_config():
    config_file = get_config_file_path()
    if os.path.exists(config_file):
        try:
            with open(config_file, 'r') as f:
                config = json.load(f)
                return config
        except json.JSONDecodeError:
            return {}
    else:
        return {}

def save_config(config):
    config_file = get_config_file_path()
    try:
        with open(config_file, 'w') as f:
            json.dump(config, f, indent=4)
    except Exception as e:
        print(f"Error al guardar la configuración: {e}")

def ensure_style_exists(doc, style_name, style_type):
    styles = doc.styles
    if style_name not in [style.name for style in styles]:
        if style_type == 'parrafo':
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        elif style_type == 'caracter':
            style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
        font = style.font
        font.size = Pt(12)
    return style_name

def apply_styles(paragraph, text, style_name, style_type, doc):
    style_name = ensure_style_exists(doc, style_name, style_type)
    if style_type == 'parrafo':
        paragraph.style = style_name
        paragraph.add_run(text)
    elif style_type == 'caracter':
        run = paragraph.add_run(text)
        run.style = style_name

def clean_default_styles(doc):
    styles = doc.styles
    keep_styles = {s['style'] for s in load_config().values()}
    keep_styles.add("Agenda-General-Parrafo")

    for style in list(styles):
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER) and style.name not in keep_styles:
            styles.element.remove(style.element)

def sanitize_xml_content(content):
    content = re.sub(r'&(?!amp;)', 'i', content)
    replacements = {}
    for old, new in replacements.items():
        content = content.replace(old, new)
    return content

def create_sanitized_copy(xml_file):
    sanitized_file = xml_file + ".sanitized"
    with open(xml_file, 'r', encoding='utf-8') as file:
        content = file.read()
    sanitized_content = sanitize_xml_content(content)
    with open(sanitized_file, 'w', encoding='utf-8') as file:
        file.write(sanitized_content)
    return sanitized_file

def validate_xml_file(xml_file):
    try:
        tree = ET.parse(xml_file)
        tree.getroot()
        return True
    except ET.ParseError:
        return False

def process_combined_elements(paragraph, text1, text2, style1, style2, doc):
    if text1:
        style1 = ensure_style_exists(doc, style1, 'caracter')
        run1 = paragraph.add_run(text1)
        run1.style = style1

    if text1 and text2:
        paragraph.add_run(" · ")

    if text2:
        style2 = ensure_style_exists(doc, style2, 'caracter')
        run2 = paragraph.add_run(text2)
        run2.style = style2

def process_fields(parent_element, fields, doc, config):
    for field in fields:
        element = parent_element.find(field)
        if element is not None and element.text and element.text.strip():
            paragraph = doc.add_paragraph()
            style_name = config.get(field, {}).get('style', "Agenda-General-Parrafo")
            style_type = config.get(field, {}).get('type', 'parrafo')
            apply_styles(paragraph, element.text.strip(), style_name, style_type, doc)

def process_activity(activity, doc, config):
    activity_title = activity.find('actividad-titulo')
    activity_time = activity.find('actividad-hora')
    activity_place = activity.find('actividad-lugar')
    activity_description = activity.find('actividad-descripcion')
    activity_extra_info = activity.find('actividad-info-extra')

    if activity_title is not None and activity_title.text:
        paragraph = doc.add_paragraph()
        apply_styles(paragraph, activity_title.text.strip(), 
                     config.get('actividad-titulo', {}).get('style', "Agenda-General-Parrafo"),
                     config.get('actividad-titulo', {}).get('type', 'parrafo'), doc)

    if activity_time is not None or activity_place is not None:
        paragraph = doc.add_paragraph()
        text1 = activity_time.text.strip() if activity_time is not None and activity_time.text else ""
        text2 = activity_place.text.strip() if activity_place is not None and activity_place.text else ""
        style1 = config.get('actividad-hora', {}).get('style', "Agenda-General-Parrafo")
        style2 = config.get('actividad-lugar', {}).get('style', "Agenda-General-Parrafo")
        process_combined_elements(paragraph, text1, text2, style1, style2, doc)

    if activity_description is not None and activity_description.text:
        paragraph = doc.add_paragraph()
        apply_styles(paragraph, activity_description.text.strip(), 
                     config.get('actividad-descripcion', {}).get('style', "Agenda-General-Parrafo"),
                     config.get('actividad-descripcion', {}).get('type', 'parrafo'), doc)
    
    if activity_extra_info is not None and activity_extra_info.text:
        paragraph = doc.add_paragraph()
        apply_styles(paragraph, activity_extra_info.text.strip(), 
                     config.get('actividad-info-extra', {}).get('style', "Agenda-General-Parrafo"),
                     config.get('actividad-info-extra', {}).get('type', 'parrafo'), doc)

def process_sub_event(sub_event, doc, config):
    sub_event_title = sub_event.find('Sub-evento-Titulo')
    sub_event_day = sub_event.find('Sub-evento-Dia')
    sub_event_time = sub_event.find('Sub-evento-Hora')
    sub_event_place = sub_event.find('Sub-evento-Lugar')
    sub_event_description = sub_event.find('Sub-evento-descripcion')
    sub_event_extra_info = sub_event.find('Sub-evento-info-extra')
    sub_event_activities = sub_event.find('Sub-evento-actividades')
    
    if sub_event_title is not None and sub_event_title.text:
        paragraph = doc.add_paragraph()
        apply_styles(paragraph, sub_event_title.text.strip(), 
                     config.get('Sub-evento-Titulo', {}).get('style', "Agenda-General-Parrafo"),
                     config.get('Sub-evento-Titulo', {}).get('type', 'parrafo'), doc)

    if sub_event_day is not None and sub_event_day.text:
        paragraph = doc.add_paragraph()
        apply_styles(paragraph, sub_event_day.text.strip(), 
                     config.get('Sub-evento-Dia', {}).get('style', "Agenda-General-Parrafo"),
                     config.get('Sub-evento-Dia', {}).get('type', 'parrafo'), doc)

    if sub_event_time is not None or sub_event_place is not None:
        paragraph = doc.add_paragraph()
        text1 = sub_event_time.text.strip() if sub_event_time is not None and sub_event_time.text else ""
        text2 = sub_event_place.text.strip() if sub_event_place is not None and sub_event_place.text else ""
        style1 = config.get('Sub-evento-Hora', {}).get('style', "Agenda-General-Parrafo")
        style2 = config.get('Sub-evento-Lugar', {}).get('style', "Agenda-General-Parrafo")
        process_combined_elements(paragraph, text1, text2, style1, style2, doc)

    if sub_event_description is not None and sub_event_description.text:
        paragraph = doc.add_paragraph()
        apply_styles(paragraph, sub_event_description.text.strip(), 
                     config.get('Sub-evento-descripcion', {}).get('style', "Agenda-General-Parrafo"),
                     config.get('Sub-evento-descripcion', {}).get('type', 'parrafo'), doc)

    if sub_event_extra_info is not None and sub_event_extra_info.text:
        paragraph = doc.add_paragraph()
        apply_styles(paragraph, sub_event_extra_info.text.strip(), 
                     config.get('Sub-evento-info-extra', {}).get('style', "Agenda-General-Parrafo"),
                     config.get('Sub-evento-info-extra', {}).get('type', 'parrafo'), doc)

    if sub_event_activities is not None:
        for activity in sub_event_activities.findall('actividad'):
            process_activity(activity, doc, config)

def process_event(event, doc, config):
    event_title = event.find('Evento-Principal-Titulo')
    event_day = event.find('Evento-Principal-Dia')
    event_time = event.find('Evento-Principal-Hora')
    event_place = event.find('Evento-Principal-Lugar')
    event_description = event.find('Evento-Principal-Descripcion')
    event_extra_info = event.find('Evento-Principal-info-extra')

    if event_title is not None and event_title.text:
        paragraph = doc.add_paragraph()
        apply_styles(paragraph, event_title.text.strip(), 
                     config.get('Evento-Principal-Titulo', {}).get('style', "Agenda-General-Parrafo"),
                     config.get('Evento-Principal-Titulo', {}).get('type', 'parrafo'), doc)

    if event_day is not None and event_day.text:
        paragraph = doc.add_paragraph()
        apply_styles(paragraph, event_day.text.strip(), 
                     config.get('Evento-Principal-Dia', {}).get('style', "Agenda-General-Parrafo"),
                     config.get('Evento-Principal-Dia', {}).get('type', 'parrafo'), doc)

    if event_time is not None or event_place is not None:
        paragraph = doc.add_paragraph()
        text1 = event_time.text.strip() if event_time is not None and event_time.text else ""
        text2 = event_place.text.strip() if event_place is not None and event_place.text else ""
        style1 = config.get('Evento-Principal-Hora', {}).get('style', "Agenda-General-Parrafo")
        style2 = config.get('Evento-Principal-Lugar', {}).get('style', "Agenda-General-Parrafo")
        process_combined_elements(paragraph, text1, text2, style1, style2, doc)

    if event_description is not None and event_description.text:
        paragraph = doc.add_paragraph()
        apply_styles(paragraph, event_description.text.strip(), 
                     config.get('Evento-Principal-Descripcion', {}).get('style', "Agenda-General-Parrafo"),
                     config.get('Evento-Principal-Descripcion', {}).get('type', 'parrafo'), doc)

    if event_extra_info is not None and event_extra_info.text:
        paragraph = doc.add_paragraph()
        apply_styles(paragraph, event_extra_info.text.strip(), 
                     config.get('Evento-Principal-info-extra', {}).get('style', "Agenda-General-Parrafo"),
                     config.get('Evento-Principal-info-extra', {}).get('type', 'parrafo'), doc)

def process_xml_to_docx(xml_file, output_folder, output_file_name):
    sanitized_file = create_sanitized_copy(xml_file)
    if not validate_xml_file(sanitized_file):
        print("Error: El archivo XML no está bien formateado después de la sanitización.")
        return

    config = load_config()
    tree = ET.parse(sanitized_file)
    root = tree.getroot()

    doc = Document()

    if "Agenda-General-Parrafo" not in [style.name for style in doc.styles]:
        general_style = doc.styles.add_style("Agenda-General-Parrafo", WD_STYLE_TYPE.PARAGRAPH)
        general_style.font.size = Pt(12)

    for event in root.findall('Evento-Principal'):
        process_event(event, doc, config)
        
        programa = event.find('Evento-Principal-Programa')
        if programa is not None:
            for sub_event in programa.findall('Sub-evento'):
                process_sub_event(sub_event, doc, config)

    clean_default_styles(doc)

    for paragraph in doc.paragraphs:
        if paragraph.style is None or paragraph.style.name == 'Normal':
            paragraph.style = "Agenda-General-Parrafo"

    output_path = os.path.join(output_folder, output_file_name)
    doc.save(output_path)

    if os.path.exists(sanitized_file):
        os.remove(sanitized_file)
